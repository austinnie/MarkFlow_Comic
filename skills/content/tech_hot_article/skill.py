"""
tech_hot_article - 技术热点文章生成器

功能：
  - 获取当前技术热点
  - 基于热点生成技术文章
  - 生成配图
  - 输出 Word 文档
"""

import os
import time
import json
import logging
import random
import re
import requests
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import hashlib

logger = logging.getLogger(__name__)

# 依赖检查
try:
    import feedparser
    FEEDPARSER_AVAILABLE = True
except ImportError:
    FEEDPARSER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


class TechHotArticle:
    """
    技术热点文章生成器
    """

    # ============ 热点源配置 ============
    HOT_SOURCES = {
        "github": "https://github.com/trending",
        "hacker_news": "https://hnrss.org/frontpage",
        "techcrunch": "https://techcrunch.com/feed/",
        "the_verge": "https://www.theverge.com/rss/index.xml",
        "wired": "https://www.wired.com/feed/rss",
        "arstechnica": "https://arstechnica.com/feed/",
        "zdnet": "https://www.zdnet.com/news/rss.xml",
        "venturebeat": "https://feeds.feedburner.com/venturebeat/SZYF",
        "devto": "https://dev.to/feed",
        "medium": "https://medium.com/feed/tag/technology",
    }

    # ============ 文章风格 ============
    WRITING_STYLES = [
        "专业分析型",
        "通俗科普型",
        "深度技术型",
        "行业观察型",
        "趋势预测型",
    ]

    # ============ 文章主题 ============
    TOPICS = [
        "人工智能",
        "机器学习",
        "大语言模型",
        "云计算",
        "边缘计算",
        "网络安全",
        "区块链",
        "Web3",
        "物联网",
        "5G",
        "量子计算",
        "生物科技",
        "新能源",
        "自动驾驶",
        "机器人",
    ]

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.name = "tech_hot_article"
        self.version = "1.0.0"
        self._setup_logging()
        self._setup_config()

        if not FEEDPARSER_AVAILABLE:
            logger.warning("feedparser 未安装，热点抓取功能可能受限")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx 未安装，Word 导出功能不可用")
        if not PIL_AVAILABLE:
            logger.warning("Pillow 未安装，图片生成功能不可用")

        logger.info("TechHotArticle 初始化完成")

    def _setup_logging(self):
        log_level = self.config.get("log_level", "INFO")
        logging.basicConfig(
            level=getattr(logging, log_level.upper()),
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )

    def _setup_config(self):
        defaults = {
            "output_dir": "./skills/tech_hot_article/output",
            "max_hot_items": 10,
            "article_words": 1500,
            "temperature": 0.85,
            "ollama_url": "http://localhost:11434",
            "model": "qwen2.5:7b",
            "image_width": 800,
            "image_height": 600,
        }
        for key, value in defaults.items():
            if key not in self.config:
                self.config[key] = value

        Path(self.config["output_dir"]).mkdir(parents=True, exist_ok=True)
        Path(self.config["output_dir"] + "/articles").mkdir(parents=True, exist_ok=True)
        Path(self.config["output_dir"] + "/images").mkdir(parents=True, exist_ok=True)
        Path(self.config["output_dir"] + "/word").mkdir(parents=True, exist_ok=True)

    # ==================== 热点获取 ====================

    def _fetch_hacker_news(self) -> List[Dict]:
        """抓取 Hacker News 热点"""
        try:
            import feedparser
            feed = feedparser.parse(self.HOT_SOURCES["hacker_news"])
            items = []
            for entry in feed.entries[:10]:
                items.append({
                    "title": entry.get("title", ""),
                    "link": entry.get("link", ""),
                    "summary": entry.get("summary", ""),
                    "source": "Hacker News",
                    "score": int(entry.get("score", 0)) if entry.get("score") else 0,
                })
            return items
        except Exception as e:
            logger.warning(f"抓取 Hacker News 失败: {e}")
            return []

    def _fetch_devto(self) -> List[Dict]:
        """抓取 Dev.to 热点"""
        try:
            import feedparser
            feed = feedparser.parse(self.HOT_SOURCES["devto"])
            items = []
            for entry in feed.entries[:10]:
                items.append({
                    "title": entry.get("title", ""),
                    "link": entry.get("link", ""),
                    "summary": entry.get("summary", ""),
                    "source": "Dev.to",
                    "score": 0,
                })
            return items
        except Exception as e:
            logger.warning(f"抓取 Dev.to 失败: {e}")
            return []

    def _fetch_techcrunch(self) -> List[Dict]:
        """抓取 TechCrunch 热点"""
        try:
            import feedparser
            feed = feedparser.parse(self.HOT_SOURCES["techcrunch"])
            items = []
            for entry in feed.entries[:10]:
                items.append({
                    "title": entry.get("title", ""),
                    "link": entry.get("link", ""),
                    "summary": entry.get("summary", ""),
                    "source": "TechCrunch",
                    "score": 0,
                })
            return items
        except Exception as e:
            logger.warning(f"抓取 TechCrunch 失败: {e}")
            return []

    def _fetch_github_trending(self) -> List[Dict]:
        """抓取 GitHub Trending"""
        try:
            import feedparser
            # GitHub Trending 没有标准 RSS，用模拟数据
            # 实际可以用 GitHub API
            topics = [
                "LangChain", "AutoGPT", "Stable Diffusion", "Whisper",
                "Llama", "ChatGPT", "Python", "Rust", "React", "Vue",
                "TypeScript", "Go", "Rust", "Zig", "Docker",
                "Kubernetes", "PyTorch", "TensorFlow", "Jupyter", "VS Code"
            ]
            items = []
            for topic in random.sample(topics, 5):
                items.append({
                    "title": f"GitHub Trending: {topic}",
                    "link": f"https://github.com/trending/{topic.lower()}",
                    "summary": f"{topic} 是当前 GitHub 上最热门的项目之一",
                    "source": "GitHub Trending",
                    "score": random.randint(100, 1000),
                })
            return items
        except Exception as e:
            logger.warning(f"抓取 GitHub Trending 失败: {e}")
            return []

    def get_hot_topics(self) -> List[Dict]:
        """获取所有热点"""
        all_items = []

        # 从多个源抓取
        all_items.extend(self._fetch_hacker_news())
        all_items.extend(self._fetch_devto())
        all_items.extend(self._fetch_techcrunch())
        all_items.extend(self._fetch_github_trending())

        # 如果抓取失败，使用模拟热点
        if not all_items:
            logger.warning("所有热点源抓取失败，使用模拟数据")
            all_items = self._get_mock_hot_topics()

        # 去重（按标题去重）
        seen = set()
        unique_items = []
        for item in all_items:
            key = item["title"][:30].lower()
            if key not in seen:
                seen.add(key)
                unique_items.append(item)

        # 按热度排序
        unique_items.sort(key=lambda x: x.get("score", 0), reverse=True)

        return unique_items[:self.config.get("max_hot_items", 10)]

    def _get_mock_hot_topics(self) -> List[Dict]:
        """模拟热点（备用）"""
        mock_items = [
            {
                "title": "ChatGPT 发布多模态功能，支持图像识别",
                "link": "https://example.com/1",
                "summary": "OpenAI 宣布 ChatGPT 新增多模态能力，用户可以通过图像与 AI 交互",
                "source": "TechCrunch",
                "score": 950,
            },
            {
                "title": "Google 发布 Gemini 2.0，性能超越 GPT-4",
                "link": "https://example.com/2",
                "summary": "Google 发布新一代 AI 模型 Gemini 2.0，在多项基准测试中领先",
                "source": "Wired",
                "score": 890,
            },
            {
                "title": "Meta 开源 Llama 3，70B 参数模型免费商用",
                "link": "https://example.com/3",
                "summary": "Meta 宣布 Llama 3 开源，支持免费商用，推动 AI 民主化",
                "source": "The Verge",
                "score": 850,
            },
            {
                "title": "苹果发布 Apple Intelligence，重新定义 Siri",
                "link": "https://example.com/4",
                "summary": "苹果在 WWDC 发布 Apple Intelligence，Siri 获得 AI 能力提升",
                "source": "ZDNet",
                "score": 800,
            },
            {
                "title": "量子计算突破：Google 实现量子霸权 2.0",
                "link": "https://example.com/5",
                "summary": "Google 在量子计算领域取得重大突破，算力大幅提升",
                "source": "Ars Technica",
                "score": 750,
            },
            {
                "title": "华为发布盘古大模型 5.0，AI 能力全面升级",
                "link": "https://example.com/6",
                "summary": "华为发布盘古大模型 5.0，在多个行业场景中实现突破",
                "source": "TechCrunch",
                "score": 720,
            },
            {
                "title": "Anthropic 推出 Claude 3.5，上下文超 200K",
                "link": "https://example.com/7",
                "summary": "Anthropic 发布 Claude 3.5，上下文窗口提升至 200K tokens",
                "source": "VentureBeat",
                "score": 680,
            },
            {
                "title": "Stability AI 发布 Stable Diffusion 3.5，图像质量提升",
                "link": "https://example.com/8",
                "summary": "Stability AI 发布新版本图像生成模型，支持更精细的创作",
                "source": "Dev.to",
                "score": 650,
            },
        ]
        return mock_items

    # ==================== 文章生成 ====================

    def _call_ollama(self, prompt: str, temperature: float = 0.85) -> str:
        """调用 Ollama API"""
        url = f"{self.config.get('ollama_url')}/api/generate"
        model = self.config.get("model", "qwen2.5:7b")

        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": temperature,
                "num_predict": 4096,
            }
        }

        try:
            response = requests.post(url, json=payload, timeout=300)
            response.raise_for_status()
            data = response.json()
            return data.get("response", "").strip()
        except Exception as e:
            logger.error(f"Ollama 调用失败: {e}")
            return ""

    def generate_article(self, hot_item: Dict, style: str = None) -> Dict:
        """生成文章"""
        if style is None:
            style = random.choice(self.WRITING_STYLES)

        title = hot_item.get("title", "")
        summary = hot_item.get("summary", "")
        source = hot_item.get("source", "")

        # 构建提示词
        prompt = f"""你是一位资深科技编辑，请根据以下热点信息，撰写一篇技术文章。

热点标题：{title}
热点描述：{summary}
来源：{source}

写作风格：{style}
文章字数：约 1500 字

要求：
1. 文章标题要吸引人
2. 开头引入热点话题
3. 正文分析技术背景、影响和未来趋势
4. 结尾总结观点
5. 语言专业但不晦涩
6. 结构清晰，分段落

请直接输出文章内容，包括标题和正文："""

        # 生成文章
        logger.info(f"正在生成文章，风格: {style}")
        content = self._call_ollama(prompt)

        if not content:
            logger.error("文章生成失败")
            return None

        # 提取标题和正文
        lines = content.strip().split("\n")
        article_title = lines[0].strip() if lines else title
        if article_title.startswith("#"):
            article_title = article_title[1:].strip()
        if not article_title:
            article_title = title

        body = "\n".join(lines[1:]) if len(lines) > 1 else content

        return {
            "title": article_title,
            "body": body,
            "full_content": content,
            "style": style,
            "hot_title": title,
            "hot_source": source,
        }

    # ==================== 图片生成 ====================

    def generate_image(self, title: str, topic: str = None) -> str:
        """生成配图"""
        output_dir = Path(self.config["output_dir"]) / "images"
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"article_image_{timestamp}.png"
        filepath = output_dir / filename

        if not PIL_AVAILABLE:
            # 如果没有 PIL，生成一个简单的文本图像（使用 HTML 替代）
            logger.warning("Pillow 未安装，生成简单占位图")
            return self._generate_placeholder_image(filepath, title)

        try:
            # 创建画布
            width = self.config.get("image_width", 800)
            height = self.config.get("image_height", 600)
            img = Image.new("RGB", (width, height), color=(20, 30, 50))

            # 绘制背景渐变（模拟）
            draw = ImageDraw.Draw(img)

            # 绘制一些装饰图形
            colors = [
                (50, 100, 200), (200, 50, 100), (50, 200, 100),
                (200, 150, 50), (150, 50, 200)
            ]
            for i in range(random.randint(3, 6)):
                x1 = random.randint(0, width)
                y1 = random.randint(0, height)
                x2 = x1 + random.randint(100, 300)
                y2 = y1 + random.randint(100, 300)
                color = random.choice(colors)
                draw.rectangle([x1, y1, x2, y2], outline=color, width=3)

            # 绘制一些圆
            for i in range(random.randint(5, 10)):
                cx = random.randint(0, width)
                cy = random.randint(0, height)
                r = random.randint(20, 80)
                color = random.choice(colors)
                draw.ellipse([cx - r, cy - r, cx + r, cy + r], outline=color, width=2)

            # 添加文字
            try:
                # 尝试加载中文字体
                font_paths = [
                    "C:/Windows/Fonts/simsun.ttc",
                    "C:/Windows/Fonts/msyh.ttc",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                ]
                font = None
                for fp in font_paths:
                    if Path(fp).exists():
                        font = ImageFont.truetype(fp, 24)
                        break
                if font is None:
                    font = ImageFont.load_default()

                # 绘制标题文字
                text = title[:50]
                bbox = draw.textbbox((0, 0), text, font=font)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                draw.text(
                    (width // 2 - tw // 2, height // 2 - th // 2),
                    text,
                    fill=(255, 255, 255),
                    font=font
                )

                # 绘制日期
                date_text = datetime.now().strftime("%Y-%m-%d")
                draw.text(
                    (width - 200, height - 40),
                    date_text,
                    fill=(150, 180, 200),
                    font=font
                )

            except Exception as e:
                logger.warning(f"文字渲染失败: {e}")

            # 保存图片
            img.save(filepath, "PNG")
            logger.info(f"图片已生成: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"图片生成失败: {e}")
            return self._generate_placeholder_image(filepath, title)

    def _generate_placeholder_image(self, filepath: Path, title: str) -> str:
        """生成占位图"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"Placeholder Image\n")
                f.write(f"Generated: {datetime.now().isoformat()}\n")
                f.write(f"Title: {title}\n")
            return str(filepath)
        except:
            return ""

    # ==================== Word 文档生成 ====================

    def create_word_document(self, article: Dict, image_path: str = None) -> str:
        """生成 Word 文档"""
        output_dir = Path(self.config["output_dir"]) / "word"
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[<>:"/\\|?*]', '', article["title"])[:30]
        filename = f"tech_article_{safe_title}_{timestamp}.docx"
        filepath = output_dir / filename

        if not DOCX_AVAILABLE:
            logger.warning("python-docx 未安装，生成 TXT 替代")
            return self._generate_txt_alternative(article, image_path, filepath)

        try:
            doc = Document()

            # 标题
            title = doc.add_heading(article["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 元信息
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            meta_paragraph.add_run(f"写作风格：{article.get('style', '未指定')}\n")
            meta_paragraph.add_run(f"热点来源：{article.get('hot_source', '未指定')}")

            doc.add_paragraph()

            # 配图
            if image_path and Path(image_path).exists() and image_path.endswith('.png'):
                try:
                    doc.add_picture(image_path, width=Inches(6))
                    doc.add_paragraph("图：文章配图", style="Caption")
                except Exception as e:
                    logger.warning(f"插入图片失败: {e}")

            doc.add_paragraph()

            # 文章正文
            body = article.get("body", "")
            if body:
                for paragraph in body.split("\n\n"):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.paragraph_format.first_line_indent = Inches(0.3)

            # 页脚
            doc.add_page_break()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.add_run(f"文章由 AI 自动生成 | {datetime.now().strftime('%Y-%m-%d')}")

            doc.save(str(filepath))
            logger.info(f"Word 文档已生成: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Word 文档生成失败: {e}")
            return self._generate_txt_alternative(article, image_path, filepath)

    def _generate_txt_alternative(self, article: Dict, image_path: str, filepath: Path) -> str:
        """生成 TXT 替代"""
        try:
            txt_path = filepath.with_suffix(".txt")
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write(f"  {article['title']}\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"写作风格：{article.get('style', '未指定')}\n")
                f.write(f"热点来源：{article.get('hot_source', '未指定')}\n\n")
                if image_path:
                    f.write(f"配图：{image_path}\n\n")
                f.write("=" * 60 + "\n\n")
                f.write(article.get("body", ""))
            return str(txt_path)
        except Exception as e:
            logger.error(f"TXT 替代生成失败: {e}")
            return ""

    # ==================== 执行入口 ====================

    def execute(self, **kwargs) -> Dict[str, Any]:
        """执行文章生成"""
        start_time = time.time()
        logger.info(f"执行技能: {self.name} (v{self.version})")

        try:
            # 获取热点
            logger.info("正在获取技术热点...")
            hot_items = self.get_hot_topics()

            if not hot_items:
                return {"status": "error", "error": "未获取到热点"}

            # 选择热点
            selected_index = kwargs.get("hot_index")
            if selected_index is not None and 0 <= selected_index < len(hot_items):
                hot_item = hot_items[selected_index]
            else:
                hot_item = random.choice(hot_items)

            # 获取风格
            style = kwargs.get("style")
            if style not in self.WRITING_STYLES:
                style = random.choice(self.WRITING_STYLES)

            logger.info(f"选中热点: {hot_item['title']}")
            logger.info(f"写作风格: {style}")

            # 生成文章
            article = self.generate_article(hot_item, style)
            if article is None:
                return {"status": "error", "error": "文章生成失败"}

            # 生成配图
            image_path = self.generate_image(article["title"], hot_item.get("source"))

            # 生成 Word 文档
            word_path = self.create_word_document(article, image_path)

            # 保存文章信息
            output_dir = Path(self.config["output_dir"]) / "articles"
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            info_file = output_dir / f"article_{timestamp}.json"
            with open(info_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "article": article,
                    "hot_item": hot_item,
                    "image_path": image_path,
                    "word_path": word_path,
                    "timestamp": datetime.now().isoformat(),
                }, f, ensure_ascii=False, indent=2)

            return {
                "status": "success",
                "result": {
                    "title": article["title"],
                    "hot_topic": hot_item["title"],
                    "hot_source": hot_item.get("source", ""),
                    "style": style,
                    "word_file": word_path,
                    "image_file": image_path,
                    "article_file": str(info_file),
                    "generated_at": datetime.now().isoformat(),
                },
                "metadata": {
                    "skill": self.name,
                    "version": self.version,
                }
            }

        except Exception as e:
            logger.error(f"执行失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                "status": "error",
                "error": str(e),
                "skill": self.name,
            }

    def __repr__(self):
        return f"<TechHotArticle(name={self.name}, version={self.version})>"